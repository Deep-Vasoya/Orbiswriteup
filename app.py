from flask import Flask, render_template, request, send_file, url_for
import pandas as pd
import docx
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from collections import defaultdict
from datetime import datetime
import os

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)

app.config['PROPAGATE_EXCEPTIONS'] = True

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'file' not in request.files:
            return render_template('index.html', error='No file part')
        file = request.files['file']
        if file.filename == '':
            return render_template('index.html', error='No selected file')
        if file and file.filename.endswith('.csv'):
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(filepath)
            currency = request.form['currency']
            from_price = request.form.get('from_price')
            to_price = request.form.get('to_price')

            try:
                from_price = int(from_price) if from_price else None
                to_price = int(to_price) if to_price else None
            except ValueError:
                return render_template('index.html', error='Invalid price range.')

            result = process_csv_and_generate_word(filepath, currency, from_price, to_price)
            if result.startswith("Error:"):
                return render_template('index.html', error=result)
            else:
                return render_template('index.html', word_file=os.path.basename(result))
        else:
            return render_template('index.html', error='Invalid file type. Please upload a CSV file.')
    return render_template('index.html')

def process_csv_and_generate_word(csv_filepath, currency, from_price=None, to_price=None):
    try:
        df = pd.read_csv(csv_filepath)

        if df.empty:
            return "Error: CSV file is empty."

        if not all(col in df.columns for col in ['package_price', 'airport', 'traveldate']):
            return "Error: CSV file is missing required columns."

        grouped_data = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))
        for _, row in df.iterrows():
            try:
                price = int(row['package_price'])
                if from_price is not None and price < from_price:
                    continue
                if to_price is not None and price > to_price:
                    continue

                price_str = f"@{currency}{price}pp"
                airport = row['airport'] + " Departure"
                date_obj = datetime.strptime(row['traveldate'], "%Y/%m/%d")
                day = str(date_obj.day)
                month = date_obj.strftime("%b %Y")  # Remove hyphen

                grouped_data[price_str][airport][month].append(day)
            except ValueError:
                return "Error: Invalid date format or package price in CSV file."

        doc = Document()
        first_price_paragraph = True

        for price, airports in grouped_data.items():
            para = doc.add_paragraph()
            run = para.add_run(price)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.name = 'Calibri (Body)'
            run.font.size = docx.shared.Pt(11)
            para.paragraph_format.line_spacing = 1 if first_price_paragraph else 1

            first_price_paragraph = False

            for airport, months in airports.items():
                para = doc.add_paragraph()
                run = para.add_run(airport)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 176, 80)
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.font.name = 'Calibri (Body)'
                run.font.size = docx.shared.Pt(11)
                para.paragraph_format.line_spacing = 1

                for month, days in months.items():
                    para = doc.add_paragraph(f"{month} – {', '.join(sorted(days, key=int))}")
                    run = para.runs[0]
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.name = 'Calibri (Body)'
                    run.font.size = docx.shared.Pt(11)
                    para.paragraph_format.line_spacing = 1

            doc.add_paragraph("")
            doc.paragraphs[-1].paragraph_format.line_spacing = 1

        csv_filename = os.path.basename(csv_filepath)
        word_filename = os.path.splitext(csv_filename)[0] + " Write Up.docx"
        word_filepath = os.path.join(app.config['OUTPUT_FOLDER'], word_filename)
        doc.save(word_filepath)

        return word_filepath

    except FileNotFoundError:
        return "Error: CSV file not found."
    except Exception as e:
        return f"An unexpected error occurred: {e}"

@app.route('/download/<filename>')
def download(filename):
    return send_file(os.path.join(app.config['OUTPUT_FOLDER'], filename), as_attachment=True)

if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=True)